"""
AI Consulting Assistant Platform - Report Exporter Service
컨설팅 보고서를 PDF, Word, HTML 형식으로 내보내기 위한 서비스
"""
from typing import Dict, Any, List, Optional
import io
import json
import re
from datetime import datetime

# Word Generation
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# PDF Generation
from xhtml2pdf import pisa

# HTML Generation (using Markdown conversion)
import markdown

# HTML Parsing
from bs4 import BeautifulSoup, Tag

class ReportExporter:
    """보고서 내보내기 서비스"""

    # ── 공통 색상 ──────────────────────────────────────────────
    COLOR_TITLE = RGBColor(0x1F, 0x4E, 0x79)   # #1F4E79 - 타이틀/H1
    COLOR_H2    = RGBColor(0x2E, 0x75, 0xB6)   # #2E75B6 - H2
    COLOR_H3    = RGBColor(0x40, 0x40, 0x40)   # #404040 - H3
    FONT_KO     = "맑은 고딕"                    # 한글 기본 폰트

    def __init__(self):
        pass

    # ─────────────────────────────────────────────────────────────
    #  DOCX  (컨설팅 보고서 템플릿 형식)
    # ─────────────────────────────────────────────────────────────

    def export_to_docx(self, report_data: Dict[str, Any]) -> bytes:
        """보고서를 Word(.docx)로 변환 - AI 컨설팅 보고서 템플릿 형식

        템플릿 규격:
          - 페이지: A4 (21 × 29.7cm), 여백 2.5cm 상하좌우
          - 표지:   타이틀 24pt bold #1F4E79 (중앙 정렬)
          - H1:     16pt bold #1F4E79
          - H2(소제목 ■): 13pt bold #2E75B6
          - H3(세부항목): 12pt bold #404040
          - 본문:   11pt 맑은고딕
        """
        doc = Document()

        # ── 페이지 설정 ──────────────────────────────────────────
        sec = doc.sections[0]
        sec.page_width  = Cm(21)
        sec.page_height = Cm(29.7)
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

        # ── 기본 스타일 오버라이드 (Normal) ─────────────────────
        normal_style = doc.styles['Normal']
        normal_style.font.name = self.FONT_KO
        normal_style.font.size = Pt(11)
        normal_style.element.rPr.rFonts.set(qn('w:eastAsia'), self.FONT_KO)

        title       = report_data.get("title", "컨설팅 보고서")
        generated_at = report_data.get("generated_at", datetime.now())
        if isinstance(generated_at, datetime):
            date_str = generated_at.strftime("%Y년 %m월 %d일")
        elif isinstance(generated_at, str) and len(generated_at) >= 10:
            date_str = generated_at[:10]
        else:
            date_str = datetime.now().strftime("%Y년 %m월 %d일")

        # ── 표지 생성 ────────────────────────────────────────────
        self._build_cover_page(doc, title, date_str)

        # ── 본문 섹션 렌더링 ─────────────────────────────────────
        sections  = report_data.get("sections",  [])
        chapters  = report_data.get("chapters",  [])

        if chapters:
            # EnhancedReportGenerator 형식: chapters → sections
            for chap in chapters:
                chap_title = f"{chap.get('number','')}. {chap.get('title','')}".strip('. ')
                self._add_heading(doc, chap_title, 1)
                for sub in chap.get("sections", []):
                    if isinstance(sub, dict):
                        self._add_heading(doc, sub.get("title",""), 2)
                        self._render_content(doc, sub.get("content",""))
        elif sections:
            if isinstance(sections, list):
                # 일반 sections 리스트: [{title, content}, ...]
                for sec_data in sections:
                    if not isinstance(sec_data, dict):
                        continue
                    self._add_heading(doc, sec_data.get("title",""), 1)
                    self._render_content(doc, sec_data.get("content",""))
            elif isinstance(sections, dict):
                # 중첩 dict 형식: {key: {title, subsections/content}, ...}
                for _key, sec_data in sections.items():
                    if not isinstance(sec_data, dict):
                        continue
                    self._add_heading(doc, sec_data.get("title",""), 1)
                    if "subsections" in sec_data:
                        for _sk, sub in sec_data["subsections"].items():
                            if isinstance(sub, dict):
                                self._add_heading(doc, sub.get("title",""), 2)
                                self._render_content(doc, sub.get("content",""))
                    elif "content" in sec_data:
                        self._render_content(doc, sec_data["content"])

        f = io.BytesIO()
        doc.save(f)
        return f.getvalue()

    # ─────────────────────────────────────────────────────────────
    #  내부 헬퍼
    # ─────────────────────────────────────────────────────────────

    def _set_run_style(self, run, size_pt: float, bold: bool = False,
                       color: RGBColor = None):
        """run에 폰트 스타일 적용 (한글 포함)"""
        run.font.name = self.FONT_KO
        run.font.size = Pt(size_pt)
        run.font.bold = bold
        if color:
            run.font.color.rgb = color
        # 한글 폰트 명시
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn('w:eastAsia'), self.FONT_KO)

    def _add_heading(self, doc: Document, text: str, level: int):
        """스타일이 적용된 제목 단락 추가"""
        if not text.strip():
            return
        p = doc.add_paragraph()
        run = p.add_run(text)
        if level == 1:
            self._set_run_style(run, 16, bold=True, color=self.COLOR_TITLE)
            p.paragraph_format.space_before = Pt(20)
            p.paragraph_format.space_after  = Pt(8)
        elif level == 2:
            self._set_run_style(run, 13, bold=True, color=self.COLOR_H2)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after  = Pt(4)
        elif level == 3:
            self._set_run_style(run, 12, bold=True, color=self.COLOR_H3)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(4)
        return p

    def _add_body(self, doc: Document, text: str, indent_cm: float = 0):
        """본문 단락 추가"""
        p = doc.add_paragraph()
        run = p.add_run(text)
        self._set_run_style(run, 11)
        p.paragraph_format.space_after = Pt(3)
        if indent_cm:
            p.paragraph_format.left_indent = Cm(indent_cm)
        return p

    def _add_bullet(self, doc: Document, text: str, indent_cm: float = 0.5):
        """글머리 기호 단락 추가"""
        p = doc.add_paragraph()
        run = p.add_run(text)
        self._set_run_style(run, 11)
        p.paragraph_format.left_indent  = Cm(indent_cm)
        p.paragraph_format.space_after  = Pt(2)
        return p

    def _render_content(self, doc: Document, content: str):
        """섹션 본문 텍스트를 파싱하여 단락 추가

        파싱 규칙:
          ■ 텍스트     → H2 소제목 (13pt bold #2E75B6)
          ①②③…⑩ 텍스트 → 번호 글머리 (들여쓰기)
          - 텍스트     → 글머리 기호 (들여쓰기)
          • 텍스트     → 글머리 기호 (들여쓰기)
          · 텍스트     → 들여쓰기 본문
          (공백 2+ 시작) → 들여쓰기 본문
          일반 텍스트   → 본문
        """
        if not content:
            return

        CIRCLED = "①②③④⑤⑥⑦⑧⑨⑩⑪⑫"

        for raw_line in content.splitlines():
            line = raw_line.rstrip()

            # 빈 줄 → 작은 여백
            if not line.strip():
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                continue

            stripped = line.strip()

            # ■ 소제목
            if stripped.startswith("■"):
                heading_text = stripped[1:].strip()
                self._add_heading(doc, heading_text, 2)

            # ① ~ ⑫ 번호 리스트
            elif stripped and stripped[0] in CIRCLED:
                self._add_bullet(doc, stripped, indent_cm=0.5)

            # - 또는 • 글머리 기호
            elif stripped.startswith("- ") or stripped.startswith("• "):
                self._add_bullet(doc, stripped, indent_cm=0.5)

            # · 점 글머리 (성숙도 바 등)
            elif stripped.startswith("·"):
                self._add_body(doc, "  " + stripped, indent_cm=0.5)

            # 들여쓰기 텍스트 (2개 이상 공백 시작)
            elif line.startswith("  "):
                self._add_body(doc, stripped, indent_cm=0.5)

            # 일반 본문
            else:
                self._add_body(doc, stripped)

    def _build_cover_page(self, doc: Document, title: str, date_str: str):
        """표지 페이지 생성

        레이아웃:
          (상단 여백 8줄)
          AI 인프라 도입 컨설팅   ← 24pt bold #1F4E79 중앙
          {보고서 유형}            ← 24pt bold #1F4E79 중앙
          (공백 2줄)
          {프로젝트명}             ← 14pt 중앙
          (공백 2줄)
          100K-AX Expert - AI/AX 10만 전문인력 양성 플랫폼  ← 11pt 중앙
          {날짜}                   ← 11pt 중앙
          (페이지 나누기)
        """
        # title 형식: "경영진 요약 보고서 — 프로젝트명"  or  "전체 보고서 — 프로젝트명"
        parts = title.split(" — ", 1)
        report_type_name = parts[0].strip() if parts else title
        project_name     = parts[1].strip() if len(parts) > 1 else ""

        # 상단 여백
        for _ in range(8):
            doc.add_paragraph()

        # 타이틀 1행
        p1 = doc.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = p1.add_run("AI 인프라 도입 컨설팅")
        self._set_run_style(run1, 24, bold=True, color=self.COLOR_TITLE)

        # 타이틀 2행 (보고서 유형)
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(report_type_name)
        self._set_run_style(run2, 24, bold=True, color=self.COLOR_TITLE)

        # 공백
        doc.add_paragraph()

        # 프로젝트명 (있을 경우)
        if project_name:
            pp = doc.add_paragraph()
            pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            runp = pp.add_run(project_name)
            self._set_run_style(runp, 14)

        # 공백 2줄
        doc.add_paragraph()
        doc.add_paragraph()

        # 컨설팅 도구명
        pt = doc.add_paragraph()
        pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        runt = pt.add_run("100K-AX Expert - AI/AX 10만 전문인력 양성 플랫폼")
        self._set_run_style(runt, 11)

        # 날짜
        pd = doc.add_paragraph()
        pd.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rund = pd.add_run(date_str)
        self._set_run_style(rund, 11)

        # 페이지 나누기
        doc.add_page_break()

    # ─────────────────────────────────────────────────────────────
    #  HTML
    # ─────────────────────────────────────────────────────────────

    def export_to_html(self, report_data: Dict[str, Any]) -> str:
        """보고서를 HTML 형식으로 변환"""
        title = report_data.get("title", "컨설팅 보고서")
        generated_at = report_data.get("generated_at", datetime.now())
        if isinstance(generated_at, datetime):
            generated_at = generated_at.strftime("%Y-%m-%d %H:%M:%S")

        html_content = f"""
<!DOCTYPE html>
<html lang="ko">
<head>
    <meta charset="UTF-8">
    <title>{title}</title>
    <style>
        body {{
            font-family: 'Malgun Gothic', 'Apple SD Gothic Neo', sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 900px;
            margin: 0 auto;
            padding: 40px;
        }}
        h1 {{ color: #1F4E79; font-size: 22px; border-bottom: 2px solid #1F4E79; padding-bottom: 8px; margin-top: 36px; }}
        h2 {{ color: #2E75B6; font-size: 16px; margin-top: 24px; }}
        h3 {{ color: #404040; font-size: 14px; margin-top: 18px; }}
        .meta {{ color: #7f8c8d; font-size: 0.9em; margin-bottom: 40px; }}
        .section {{ margin-bottom: 30px; }}
        table {{ width: 100%; border-collapse: collapse; margin: 20px 0; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
        th {{ background-color: #1F4E79; color: white; font-weight: bold; }}
        hr {{ border: 0; border-top: 1px solid #eee; margin: 30px 0; }}
        pre {{ white-space: pre-wrap; font-family: inherit; }}
    </style>
</head>
<body>
    <h1 style="font-size:28px; text-align:center; border:none;">{title}</h1>
    <div class="meta" style="text-align:center;">생성일: {generated_at}</div>
"""

        sections = report_data.get("sections", [])
        chapters = report_data.get("chapters", [])

        if chapters:
            for chapter in chapters:
                html_content += f"<h1>{chapter.get('number', '')}. {chapter.get('title', '')}</h1>"
                for section in chapter.get("sections", []):
                    sec_title   = section.get("title", "") if isinstance(section, dict) else str(section)
                    sec_content = section.get("content", "") if isinstance(section, dict) else ""
                    html_content += f"<h2>{sec_title}</h2>"
                    if sec_content:
                        html_content += markdown.markdown(sec_content, extensions=['tables'])

        elif sections:
            if isinstance(sections, list):
                for section in sections:
                    if isinstance(section, dict):
                        html_content += f"<div class='section'><h1>{section.get('title', '')}</h1>"
                        content = section.get('content', '')
                        # ■ → h2 변환
                        lines = []
                        for ln in content.splitlines():
                            s = ln.strip()
                            if s.startswith('■'):
                                lines.append(f"<h2>{s[1:].strip()}</h2>")
                            else:
                                lines.append(ln)
                        html_content += markdown.markdown('\n'.join(lines), extensions=['tables'])
                        html_content += "</div><hr>"
            elif isinstance(sections, dict):
                for _key, sec_data in sections.items():
                    if not isinstance(sec_data, dict):
                        continue
                    html_content += f"<div class='section'><h1>{sec_data.get('title','')}</h1>"
                    if "subsections" in sec_data:
                        for _sk, sub in sec_data["subsections"].items():
                            if isinstance(sub, dict):
                                html_content += f"<h2>{sub.get('title','')}</h2>"
                                html_content += markdown.markdown(
                                    sub.get("content",""), extensions=['tables']
                                )
                    elif "content" in sec_data:
                        html_content += markdown.markdown(
                            sec_data["content"], extensions=['tables']
                        )
                    html_content += "</div><hr>"

        html_content += """
</body>
</html>
"""
        return html_content

    # ─────────────────────────────────────────────────────────────
    #  PDF
    # ─────────────────────────────────────────────────────────────

    def export_to_pdf(self, report_data: Dict[str, Any]) -> bytes:
        """보고서를 PDF로 변환 (Bytes 반환)"""
        from pathlib import Path

        # 한글 폰트 경로 (NanumGothic TTF)
        fonts_dir = Path(__file__).resolve().parent.parent.parent / "fonts"
        font_regular = str(fonts_dir / "NanumGothic-Regular.ttf")
        font_bold = str(fonts_dir / "NanumGothic-Bold.ttf")

        source_html = self.export_to_html(report_data)

        font_style = f"""
        <style>
            @font-face {{
                font-family: NanumGothic;
                src: url({font_regular});
            }}
            @font-face {{
                font-family: NanumGothic;
                font-weight: bold;
                src: url({font_bold});
            }}
            @page {{ size: A4; margin: 2.5cm; }}
            * {{ font-family: NanumGothic, sans-serif; }}
        </style>
        """
        source_html = source_html.replace("</head>", f"{font_style}</head>")

        result_file = io.BytesIO()
        pisa_status = pisa.CreatePDF(
            source_html,
            dest=result_file,
            encoding='UTF-8'
        )

        if pisa_status.err:
            raise Exception(f"PDF generation error: {pisa_status.err}")

        return result_file.getvalue()

    # ─────────────────────────────────────────────────────────────
    #  HTML → PDF  (ISO 24030 미리보기 HTML 직접 변환)
    # ─────────────────────────────────────────────────────────────

    def export_html_to_pdf(self, html_content: str) -> bytes:
        """ISO 24030 보고서 미리보기 HTML을 PDF로 변환

        Args:
            html_content: 브라우저 innerHTML (Bootstrap 스타일 포함)
        Returns:
            PDF 바이트
        """
        from pathlib import Path

        fonts_dir = Path(__file__).resolve().parent.parent.parent / "fonts"
        font_regular = str(fonts_dir / "NanumGothic-Regular.ttf")
        font_bold    = str(fonts_dir / "NanumGothic-Bold.ttf")

        full_html = f"""<!DOCTYPE html>
<html lang="ko">
<head>
<meta charset="UTF-8">
<style>
    @font-face {{
        font-family: NanumGothic;
        src: url({font_regular});
    }}
    @font-face {{
        font-family: NanumGothic;
        font-weight: bold;
        src: url({font_bold});
    }}
    @page {{ size: A4; margin: 2cm; }}
    * {{ font-family: NanumGothic, 'Noto Sans KR', sans-serif; }}
    body {{ padding: 0; margin: 0; font-size: 10px; line-height: 1.5; color: #333; }}
    h3 {{ font-size: 18px; color: #1F4E79; margin-top: 16px; }}
    h5 {{ font-size: 13px; color: #2E75B6; margin-top: 12px; }}
    h6 {{ font-size: 11px; color: #404040; }}
    .badge {{ padding: 2px 6px; border-radius: 4px; font-size: 9px; color: white; }}
    .bg-primary {{ background-color: #0d6efd; }}
    .bg-success {{ background-color: #198754; }}
    .bg-warning {{ background-color: #ffc107; color: #333 !important; }}
    .bg-info {{ background-color: #0dcaf0; }}
    .bg-danger {{ background-color: #dc3545; }}
    .bg-secondary {{ background-color: #6c757d; }}
    .text-white {{ color: white; }}
    .text-dark {{ color: #212529; }}
    .text-muted {{ color: #6c757d; }}
    .fw-bold {{ font-weight: bold; }}
    .fs-4 {{ font-size: 16px; }}
    .fs-6 {{ font-size: 11px; }}
    table {{ width: 100%; border-collapse: collapse; margin: 8px 0; font-size: 9px; }}
    th, td {{ border: 1px solid #dee2e6; padding: 5px 8px; text-align: left; }}
    th {{ background-color: #1F4E79; color: white; font-weight: bold; }}
    .card {{ border: 1px solid #dee2e6; border-radius: 6px; margin-bottom: 8px; padding: 8px; }}
    .card-body {{ padding: 6px; }}
    .progress {{ background-color: #e9ecef; border-radius: 4px; height: 8px; overflow: hidden; }}
    .progress-bar {{ height: 100%; }}
    .border-bottom {{ border-bottom: 1px solid #dee2e6; }}
    .mb-4 {{ margin-bottom: 12px; }}
    .mb-3 {{ margin-bottom: 10px; }}
    .mb-2 {{ margin-bottom: 6px; }}
    .mb-1 {{ margin-bottom: 3px; }}
    .mt-2 {{ margin-top: 6px; }}
    .pb-2 {{ padding-bottom: 6px; }}
    .pb-3 {{ padding-bottom: 10px; }}
    .p-4 {{ padding: 12px; }}
    .p-3 {{ padding: 10px; }}
    .py-3 {{ padding-top: 10px; padding-bottom: 10px; }}
    .px-3 {{ padding-left: 10px; padding-right: 10px; }}
    .text-center {{ text-align: center; }}
    .rpt-section-title {{ font-size: 12px; font-weight: bold; border-bottom: 1px solid #dee2e6; padding-bottom: 5px; margin-bottom: 8px; }}
    .shadow-sm {{ /* skip shadow for PDF */ }}
    .rounded {{ border-radius: 6px; }}
    .border {{ border: 1px solid #dee2e6; }}
    .border-0 {{ border: none; }}
    .row {{ width: 100%; }}
    .d-block {{ display: block; }}
    .me-2 {{ margin-right: 6px; }}
    small {{ font-size: 9px; }}
</style>
</head>
<body>
{html_content}
</body>
</html>"""

        result_file = io.BytesIO()
        pisa_status = pisa.CreatePDF(
            full_html,
            dest=result_file,
            encoding='UTF-8'
        )

        if pisa_status.err:
            raise Exception(f"HTML→PDF 변환 오류: {pisa_status.err}")

        return result_file.getvalue()

    # ─────────────────────────────────────────────────────────────
    #  HTML → DOCX  (ISO 24030 미리보기 HTML → Word 문서)
    # ─────────────────────────────────────────────────────────────

    def export_html_to_docx(self, html_content: str, title: str = "AI 평가 보고서") -> bytes:
        """ISO 24030 보고서 미리보기 HTML을 Word 문서로 변환

        Args:
            html_content: 브라우저 innerHTML
            title: 문서 타이틀
        Returns:
            DOCX 바이트
        """
        soup = BeautifulSoup(html_content, "html.parser")
        doc = Document()

        # ── 페이지 설정 ──────────────────────────────────────────
        sec = doc.sections[0]
        sec.page_width  = Cm(21)
        sec.page_height = Cm(29.7)
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

        # ── 기본 스타일 ──────────────────────────────────────────
        normal_style = doc.styles['Normal']
        normal_style.font.name = self.FONT_KO
        normal_style.font.size = Pt(11)
        normal_style.element.rPr.rFonts.set(qn('w:eastAsia'), self.FONT_KO)

        # ── 표지 ─────────────────────────────────────────────────
        for _ in range(6):
            doc.add_paragraph()
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_title.add_run("ISO 24030 AI 평가")
        self._set_run_style(run, 24, bold=True, color=self.COLOR_TITLE)

        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p_sub.add_run(title)
        self._set_run_style(run2, 20, bold=True, color=self.COLOR_TITLE)

        doc.add_paragraph()
        p_date = doc.add_paragraph()
        p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_d = p_date.add_run(datetime.now().strftime("%Y년 %m월 %d일"))
        self._set_run_style(run_d, 11)

        doc.add_page_break()

        # ── HTML 요소 순회하며 DOCX로 변환 ─────────────────────
        self._walk_html_to_docx(doc, soup)

        f = io.BytesIO()
        doc.save(f)
        return f.getvalue()

    def _walk_html_to_docx(self, doc: Document, element):
        """HTML 요소를 재귀적으로 순회하며 DOCX 단락/테이블로 변환"""
        for child in element.children:
            if not isinstance(child, Tag):
                # 텍스트 노드 — 의미 있는 텍스트만 추가
                text = child.strip() if hasattr(child, 'strip') else str(child).strip()
                if text:
                    self._add_body(doc, text)
                continue

            tag = child.name.lower() if child.name else ""

            # 헤딩 태그
            if tag in ("h1", "h2", "h3"):
                text = child.get_text(strip=True)
                level_map = {"h1": 1, "h2": 2, "h3": 3}
                self._add_heading(doc, text, level_map[tag])

            elif tag in ("h4", "h5", "h6"):
                text = child.get_text(strip=True)
                self._add_heading(doc, text, 3)

            # 테이블
            elif tag == "table":
                self._html_table_to_docx(doc, child)

            # KPI 카드 행 (row with cards)
            elif tag == "div" and "row" in child.get("class", []):
                cards = child.find_all("div", class_="card")
                if cards:
                    self._html_cards_to_docx_table(doc, cards)
                else:
                    self._walk_html_to_docx(doc, child)

            # 보고서 섹션 (.rpt-section)
            elif tag == "div" and "rpt-section" in " ".join(child.get("class", [])):
                self._walk_html_to_docx(doc, child)

            # 보고서 섹션 제목
            elif tag in ("h5",) and "rpt-section-title" in " ".join(child.get("class", [])):
                text = child.get_text(strip=True)
                # 숫자 배지 제거 (예: "1 항목명" → "항목명")
                text = re.sub(r'^\d+\s*', '', text)
                self._add_heading(doc, text, 2)

            # 리스트
            elif tag in ("ul", "ol"):
                for li in child.find_all("li", recursive=False):
                    self._add_bullet(doc, "• " + li.get_text(strip=True))

            # 단락
            elif tag == "p":
                text = child.get_text(strip=True)
                if text:
                    self._add_body(doc, text)

            # 기타 div / section — 재귀
            elif tag in ("div", "section", "span", "article", "main", "body"):
                self._walk_html_to_docx(doc, child)

            # small, strong, em 등 인라인 요소
            elif tag in ("small", "strong", "em", "b", "i"):
                text = child.get_text(strip=True)
                if text:
                    self._add_body(doc, text)

    def _html_table_to_docx(self, doc: Document, table_el):
        """HTML <table>을 DOCX 테이블로 변환"""
        rows = table_el.find_all("tr")
        if not rows:
            return

        # 열 수 계산
        first_row_cells = rows[0].find_all(["th", "td"])
        col_count = len(first_row_cells)
        if col_count == 0:
            return

        tbl = doc.add_table(rows=0, cols=col_count)
        tbl.style = 'Table Grid'
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        for row_el in rows:
            cells = row_el.find_all(["th", "td"])
            if not cells:
                continue
            row = tbl.add_row()
            for i, cell_el in enumerate(cells):
                if i >= col_count:
                    break
                text = cell_el.get_text(strip=True)
                cell = row.cells[i]
                cell.text = ""
                p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                run = p.add_run(text)
                is_header = cell_el.name == "th"
                self._set_run_style(run, 9, bold=is_header)
                if is_header:
                    from docx.oxml import OxmlElement
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), '1F4E79')
                    shading.set(qn('w:val'), 'clear')
                    cell._element.get_or_add_tcPr().append(shading)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        doc.add_paragraph()  # 테이블 후 여백

    def _html_cards_to_docx_table(self, doc: Document, cards):
        """KPI 카드들을 DOCX 테이블 한 행으로 변환"""
        col_count = len(cards)
        if col_count == 0:
            return

        tbl = doc.add_table(rows=1, cols=col_count)
        tbl.style = 'Table Grid'
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, card in enumerate(cards):
            cell = tbl.rows[0].cells[i]
            cell.text = ""
            texts = [t.strip() for t in card.stripped_strings]
            p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if texts:
                run_val = p.add_run(texts[0])
                self._set_run_style(run_val, 14, bold=True, color=self.COLOR_TITLE)
            for t in texts[1:]:
                p2 = cell.add_paragraph()
                p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_lbl = p2.add_run(t)
                self._set_run_style(run_lbl, 9)

        doc.add_paragraph()


# Singleton
_report_exporter = None

def get_report_exporter() -> ReportExporter:
    global _report_exporter
    if _report_exporter is None:
        _report_exporter = ReportExporter()
    return _report_exporter
